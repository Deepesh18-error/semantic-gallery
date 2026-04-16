import os
import base64
import io
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
from .database import db
from .file_splitter import split_media
import subprocess
import logging
import time
from datetime import datetime

logger = logging.getLogger(__name__)


# ============================================================
# THUMBNAIL GENERATOR (VIDEO)
# ============================================================
def generate_video_thumbnail(video_path, media_id):
    start = time.time()
    print(f"    ⏱️  [PREPROCESS] VIDEO THUMBNAIL: FFmpeg extraction starting...")
    
    from django.conf import settings
    thumb_dir = os.path.join(settings.MEDIA_VAULT, "thumbnails")
    os.makedirs(thumb_dir, exist_ok=True)
    thumb_path = os.path.join(thumb_dir, f"{media_id}.jpg")

    # Normalize slashes for FFmpeg on Windows
    vp = video_path.replace('\\', '/')
    tp = thumb_path.replace('\\', '/')

    # Use -ss 0 to grab frame 0 — safest for all video lengths
    cmd = f'ffmpeg -y -ss 0 -i "{vp}" -vframes 1 -q:v 2 "{tp}"'

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, shell=True)
        if result.returncode == 0 and os.path.exists(thumb_path):
            print(f"✅ Video Thumbnail Success: {thumb_path}")
            print(f"    ⏱️  [PREPROCESS] VIDEO THUMBNAIL done in {time.time() - start:.3f}s")
            return thumb_path
        else:
            print(f"❌ FFmpeg Failed (code={result.returncode}): {result.stderr}")
            return None
    except Exception as e:
        print(f"❌ Subprocess Crash: {str(e)}")
        return None


# ============================================================
# THUMBNAIL GENERATOR (PDF) — Uses PyMuPDF, already installed
# ============================================================
def generate_pdf_thumbnail(pdf_path, media_id):
    start = time.time()
    print(f"    ⏱️  [PREPROCESS] PDF THUMBNAIL: PyMuPDF rendering page 1...")
    from django.conf import settings
    thumb_dir = os.path.join(settings.MEDIA_VAULT, "thumbnails")
    os.makedirs(thumb_dir, exist_ok=True)
    thumb_path = os.path.join(thumb_dir, f"{media_id}.jpg")

    try:
        doc = fitz.open(pdf_path)
        if doc.page_count == 0:
            print(f"❌ PDF has no pages: {pdf_path}")
            return None

        # Render first page at 2x zoom for sharp thumbnail
        page = doc[0]
        mat = fitz.Matrix(2.0, 2.0)  # 2x zoom = ~150 DPI equivalent
        pix = page.get_pixmap(matrix=mat)

        # Convert to JPEG via Pillow for quality control
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img.thumbnail((800, 1000), Image.Resampling.LANCZOS)  # Cap size
        img.save(thumb_path, "JPEG", quality=85)

        doc.close()
        print(f"✅ PDF Thumbnail Success: {thumb_path}")
        print(f"    ⏱️  [PREPROCESS] PDF THUMBNAIL done in {time.time() - start:.3f}s")
        return thumb_path

    except Exception as e:
        print(f"❌ PDF Thumbnail Failed: {str(e)}")
        return None


def generate_text_thumbnail(path, media_id, ext):
    start = time.time()
    print(f"    ⏱️  [PREPROCESS] TEXT THUMBNAIL: rendering text preview...")

    from django.conf import settings
    thumb_dir = os.path.join(settings.MEDIA_VAULT, "thumbnails")
    os.makedirs(thumb_dir, exist_ok=True)
    thumb_path = os.path.join(thumb_dir, f"{media_id}.jpg")

    try:
        # Step 1: Extract first ~500 chars of text
        preview_text = ""

        if ext == '.txt' or ext == '.md':
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                preview_text = f.read(500)

        elif ext == '.docx':
            doc = DocxDocument(path)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            preview_text = "\n".join(lines)[:500]

        if not preview_text.strip():
            preview_text = f"[Empty file: {os.path.basename(path)}]"

        # Step 2: Render text onto a white image using Pillow
        img_width, img_height = 800, 1000
        background_color = (255, 255, 255)
        text_color = (30, 30, 30)
        header_color = (99, 102, 241)  # Brand indigo
        line_height = 28
        padding = 48
        font_size = 22

        img = Image.new("RGB", (img_width, img_height), color=background_color)

        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)

        # Try to load a monospace font, fall back to default
        try:
            font = ImageFont.truetype("cour.ttf", font_size)       # Windows Courier
            header_font = ImageFont.truetype("courbd.ttf", font_size + 2)
        except:
            try:
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf", font_size)
                header_font = font
            except:
                font = ImageFont.load_default()
                header_font = font

        # Draw header bar
        draw.rectangle([(0, 0), (img_width, 64)], fill=header_color)
        ext_label = ext.upper().replace('.', '')
        draw.text((padding, 18), f"{ext_label} FILE", fill=(255, 255, 255), font=header_font)

        # Draw filename
        filename = os.path.basename(path)
        draw.text((padding, 80), filename[:60], fill=(120, 120, 120), font=font)

        # Draw divider line
        draw.line([(padding, 116), (img_width - padding, 116)], fill=(220, 220, 220), width=1)

        # Draw text lines
        y = 130
        for line in preview_text.splitlines():
            if y + line_height > img_height - padding:
                draw.text((padding, y), "...", fill=(160, 160, 160), font=font)
                break
            # Truncate long lines
            if len(line) > 60:
                line = line[:57] + "..."
            draw.text((padding, y), line, fill=text_color, font=font)
            y += line_height

        img.save(thumb_path, "JPEG", quality=85)
        print(f"✅ Text Thumbnail Success: {thumb_path}")
        print(f"    ⏱️  [PREPROCESS] TEXT THUMBNAIL done in {time.time() - start:.3f}s")
        return thumb_path

    except Exception as e:
        print(f"❌ Text Thumbnail Failed: {str(e)}")
        return None



# ============================================================
# MAIN DISPATCHER — THIS WAS THE BROKEN FUNCTION
# ============================================================
def process_media_item(media_id):
    item = db.media_items.find_one({"_id": media_id})
    if not item:
        return []

    m_type = item['media_type']
    path = item['file_path']

    if m_type == 'IMAGE':
        return _process_image(path)

    elif m_type == 'VIDEO':
        # Generate thumbnail first, save path to DB
        thumb_path = generate_video_thumbnail(path, media_id)
        if thumb_path:
            db.media_items.update_one(
                {"_id": media_id},
                {"$set": {"file_metadata.thumbnail_path": thumb_path}}
            )
        return _process_multimedia(path, media_id, is_video=True)

    elif m_type == 'AUDIO':
        return _process_multimedia(path, media_id, is_video=False)

    elif m_type == 'DOCUMENT':
        ext = os.path.splitext(path)[1].lower()

        thumb_path = None

        if ext == '.pdf':
            # PDF → render page 1 as image
            thumb_path = generate_pdf_thumbnail(path, media_id)

        elif ext in ['.txt', '.md', '.docx']:
            # Text/Docx → render text content as styled image
            thumb_path = generate_text_thumbnail(path, media_id, ext)

        # Save thumbnail path to DB for any type that generated one
        if thumb_path:
            db.media_items.update_one(
                {"_id": media_id},
                {"$set": {"file_metadata.thumbnail_path": thumb_path}}
            )

        return _process_document(path)

    # TEXT type (plain text notes typed directly)
    elif m_type == 'TEXT':
        return _process_text_content(item.get('text_content', ''))

    return []


# ============================================================
# IMAGE PROCESSOR
# ============================================================
def _process_image(path):
    start = time.time()
    print(f"    ⏱️  [PREPROCESS] IMAGE: reading + resizing + base64 encoding...")
    with Image.open(path) as img:
        if img.mode != 'RGB':
            img = img.convert('RGB')
        img.thumbnail((1024, 1024), Image.Resampling.LANCZOS)
        buffered = io.BytesIO()
        img.save(buffered, format="JPEG", quality=85)
        b64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

    print(f"    ⏱️  [PREPROCESS] IMAGE done in {time.time() - start:.3f}s | b64 size: {len(b64)/1024:.1f} KB")
    return [{"type": "IMAGE", "data": b64, "metadata": {"chunk_index": 0}}]


# ============================================================
# DOCUMENT PROCESSOR — Fixed edge cases
# ============================================================
def _process_document(path):
    import time
    start = time.time()
    ext = os.path.splitext(path)[1].lower()
    print(f"    ⏱️  [PREPROCESS] DOCUMENT ({ext}): extracting text...")

    # Step 1: Extract text — ALL variables initialized before try block
    text_content = ""

    try:
        if ext == '.pdf':
            doc = fitz.open(path)
            text_content = " ".join([page.get_text() for page in doc])
            doc.close()

        elif ext == '.docx':
            doc = DocxDocument(path)
            text_content = " ".join([
                p.text for p in doc.paragraphs if p.text.strip()
            ])

        elif ext in ['.txt', '.md']:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()

        else:
            # Fallback — try reading as plain text
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()

    except Exception as e:
        print(f"    ❌ [PREPROCESS] Text extraction crashed: {str(e)}")
        # Return placeholder so pipeline doesn't fail completely
        return [{
            "type": "TEXT",
            "data": f"Document: {os.path.basename(path)}",
            "metadata": {
                "chunk_index": 0,
                "chunk_text": "Text extraction failed for this document."
            }
        }]

    # Step 2: Clean and validate — OUTSIDE try block, text_content guaranteed to exist
    text_content = text_content.strip()

    extract_time = time.time() - start
    
    if not text_content:
        print(f"    ⚠️  [PREPROCESS] Empty document — no extractable text ({extract_time:.3f}s)")
        return [{
            "type": "TEXT",
            "data": f"Document file: {os.path.basename(path)}",
            "metadata": {
                "chunk_index": 0,
                "chunk_text": "No text content extractable from this document."
            }
        }]

    # Step 3: Chunking — words defined here, guaranteed text_content exists
    words = text_content.split()
    print(f"    ⏱️  [PREPROCESS] Text extracted in {extract_time:.3f}s | Words: {len(words)}")

    chunk_size = 500
    overlap = 50
    chunks = []
    chunk_start = time.time()

    i = 0
    while i < len(words):
        chunk_words = words[i : i + chunk_size]
        chunk_text = " ".join(chunk_words)

        chunks.append({
            "type": "TEXT",
            "data": chunk_text,
            "metadata": {
                "chunk_index": len(chunks),
                "chunk_text": chunk_text[:200]
            }
        })

        # If this chunk reached the end of words, stop
        if i + chunk_size >= len(words):
            break

        i += (chunk_size - overlap)

    print(f"    ⏱️  [PREPROCESS] Chunking done in {time.time() - chunk_start:.3f}s | Chunks: {len(chunks)}")
    return chunks



# ============================================================
# TEXT CONTENT PROCESSOR (for plain text notes)
# ============================================================
def _process_text_content(text_content):
    if not text_content or not text_content.strip():
        return []

    return [{
        "type": "TEXT",
        "data": text_content,
        "metadata": {
            "chunk_index": 0,
            "chunk_text": text_content[:200]
        }
    }]


# ============================================================
# MULTIMEDIA PROCESSOR (VIDEO / AUDIO segments)
# ============================================================
def _process_multimedia(path, media_id, is_video):
    segments = split_media(path, media_id, is_video)
    packages = []

    for seg in segments:
        packages.append({
            "type": "VIDEO" if is_video else "AUDIO",
            "file_path": seg['path'],
            "metadata": {
                "segment_index": seg['index'],
                "start_time": seg['start_time'],
                "end_time": seg['end_time']
            }
        })
    return packages